"""docx 读取 — 段落 + 表格 (Markdown 风格), 不抽样式/批注/嵌入对象.

为什么不抓样式: agent 拿到样式只能描述, 没决策价值. 同样不抓 inline 图片
(LLM 看不见二进制), agent 真要看图请用 Read 工具走 SDK.
"""

from __future__ import annotations

from pathlib import Path

from pentaloom.capabilities.file._models import FileReadResult, FileWarning


def _table_to_markdown(table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([(cell.text or "").strip().replace("\n", " ") for cell in row.cells])
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    lines = ["| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def read_docx(path: Path) -> FileReadResult:
    # 延迟 import: 装 python-docx 慢, 避免无 docx 任务时也付出 import 成本
    from docx import Document  # type: ignore[import-not-found]

    doc = Document(str(path))
    parts: list[str] = []
    warnings: list[FileWarning] = []

    # body 段落和表格混排, 按文档原顺序输出. python-docx 不直接给"按出现顺序"的
    # iterator, 这里走 body.iter_inner_content() — 1.0+ 才有, 旧版本回退到 paragraphs+tables.
    iter_method = getattr(doc, "iter_inner_content", None)
    if iter_method is None:
        for p in doc.paragraphs:
            t = (p.text or "").rstrip()
            if t:
                parts.append(t)
        for tbl in doc.tables:
            md = _table_to_markdown(tbl)
            if md:
                parts.append("\n" + md + "\n")
        warnings.append(FileWarning(
            kind="docx_unordered",
            message="python-docx 旧版本: 段落与表格未按原始顺序混排",
        ))
    else:
        for block in iter_method():
            if hasattr(block, "text"):
                t = (block.text or "").rstrip()
                if t:
                    parts.append(t)
            elif hasattr(block, "rows"):
                md = _table_to_markdown(block)
                if md:
                    parts.append("\n" + md + "\n")

    text = "\n\n".join(parts).strip()
    meta: dict[str, int | str | list[str]] = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "chars": len(text),
    }

    if not text:
        warnings.append(FileWarning(kind="empty", message="docx 没有可读文本"))

    return FileReadResult(
        path=str(path),
        kind="docx",
        text=text,
        warnings=warnings,
        meta=meta,
    )
